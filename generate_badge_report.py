'''
This script generates a report of which badges each Scout has.
'''

import os
import sys

from datetime import date
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from osm import AwardScheme, Connection, Manager, BadgeOrder

class ReportGenerator(object):

    def __init__(self):
        self._conn = None
        self._mgr = None
        self._section = None
        self._term = None
        self._badge_order = None

    def _connect(self):
        self._conn = Connection('secret.json')
        self._conn.connect()

    def _initialise(self):
        self._mgr = Manager()
        self._mgr.load(self._conn)

    def run(self):
        if len(sys.argv) < 3:
            print('ERROR: term and section have not been set! ')
            return

        print('Connecting to OSM...')
        self._connect()
        self._initialise()

        self._set_term(sys.argv[1:3])
        if self._term is None:
            return

        print('Retrieving badge order...')
        self._badge_order = BadgeOrder('report-order.json')

        print('Retrieving badge report...')
        report = self._term.load_badges_by_person(self._conn)

        print('Generating report...')
        filename = ensureExtension(sys.argv[2]+'-Badge Report', '.docx')
        document = Document()
        self._generate_report(report, document)
        self._badge_order.save('report-order.json')

        print('Saving to %s...' % (filename, ))
        document.save(filename)

        print('Done')

    def _set_term(self, args):
        term_name = args[0]
        self._set_section(args[1:])

        print('Setting term...')
        if term_name == 'current':
            term = self._section.current_term()
            if term is None:
                print('-> Currently not in a term')
                return
            else:
                self._term = term
                print('-> Term set to %s' % (str(term), ))
                return
        
        for term in self._section.terms:
            if term.name == term_name:
                self._term = term
                print('-> Term set to %s' % (str(term), ))
                return
            
        print('-> Unknown term: %s' % (term_name, ))

    def _set_section(self, args):
        print('Setting section...')
        section = self._mgr.find_section(args[0])
        if section is None:
            print('-> Unknown section: %s' % (args[0], ))
        else:
            self._section = section
            print('-> Section set to %s' % (str(section), ))
    
    def _generate_report(self, report, document):
        headingStyle = document.styles.add_style('TableHeading', WD_STYLE_TYPE.PARAGRAPH)
        headingStyle.font.bold=True

        section = document.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        section.header.paragraphs[0].text = 'Badge Report'
        now = date.today()
        section.footer.paragraphs[0].text = 'Generated ' + now.strftime('%d %B %Y')
        table = document.add_table(rows = 1, cols = 2, style='Table Grid')
        table.columns[0].width = Cm(5)
        table.columns[1].width = Cm(21)
        cells = table.rows[0].cells
        cells[0].text = 'Person'
        cells[1].text = 'Badges'
        cells[0].width = Cm(5)
        cells[1].width = Cm(21)
        clearFormatting(cells[0].paragraphs[0], headingStyle)
        clearFormatting(cells[1].paragraphs[0], headingStyle)
        set_repeat_table_header(table.rows[0])
        for person in report:
            cells = table.add_row().cells
            name = '%s %s' % (person.first_name, person.last_name)
            print('...adding row for %s...' % (name, ))
            cells[0].text = name
            clearFormatting(cells[0].paragraphs[0])
            para = cells[1].paragraphs[0]
            clearFormatting(para)
            cells[0].width = Cm(4)
            cells[1].width = Cm(22)
            person.badges.sort(key=self._sort_order)
            all_badges = { b.badge_id : True for b in person.badges if b.completed }
            for badge in person.badges:
                if self._badge_order.remove_with(badge.badge_id) in all_badges:
                    continue

                _, file_extension = os.path.splitext(badge.picture)
                badge_path = os.path.join('badge_images', badge.name + file_extension)
                if not os.path.exists(badge_path):
                    print('...retrieving badge image for %s...' % (badge.name,))
                    self._conn.download_binary(badge.picture, badge_path)
                if badge.completed:
                    para.add_run().add_picture(badge_path, width = Cm(2))
                    para.add_run(' ')
        preventDocumentBreak(document)

    def _sort_order(self, badge):
        return self._badge_order.get_order(badge.badge_id, badge.name)

def ensureExtension(filename, extension):
    return filename if filename.lower().endswith(extension) else filename + extension

def clearFormatting(paragraph, style=None):
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    if style is not None:
        paragraph.style = style

def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

def preventDocumentBreak(document):
  tags = document.element.xpath('//w:tr')
  rows = len(tags)
  for row in range(0,rows):
    tag = tags[row]
    child = OxmlElement('w:cantSplit')
    tag.append(child)

if __name__ == "__main__":
    mgr = ReportGenerator()
    mgr.run()
