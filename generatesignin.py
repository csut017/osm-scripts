''' Script for generating the weekly sign-in sheets. '''

from osm import Connection, Manager
import docx

def main():
    ''' Main script. '''
    conn = connect()
    mgr = initialise(conn)
    generate_signin(mgr, conn)


def connect():
    print('Connecting')
    conn = Connection('secret.json')
    conn.connect()
    return conn


def initialise(conn):
    print('Downloading sections and terms')
    mgr = Manager()
    mgr.load(conn)
    return mgr


def generate_signin(mgr, conn):
    section = mgr.find_section(name='Keas')
    print('Downloading members for ' + str(section))
    term = section.terms[-1]
    term.load_members(conn)
    for member in term.members:
        print('   %s' % (str(member)))

    print('Gnerating sign-in sheet')
    doc = docx.Document()
    doc.add_paragraph('Kea Members')
    table = doc.add_table(rows=1, cols=6)
    fill_row(table, 0, '', 'Name', 'Sign In', 'Sign Out', 'Contact Person', 'Contact Number')
    doc.save('signin.docx')

def fill_row(table, row_number, *values):
    if row_number >= len(table.rows):
        table.add_row()
    row = table.rows[row_number]
    cell = 0
    for value in values:
        row.cells[cell].text = value
        cell += 1


if __name__ == "__main__":
    main()
