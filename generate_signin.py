'''
This script generates a generic sign-in sheet.
'''

import os
import sys

from datetime import date
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt
from osm import AwardScheme, Connection, Manager


class ReportGenerator(object):

    def __init__(self):
        self._conn = None
        self._mgr = None
        self._section = None
        self._term = None

    def _connect(self):
        self._conn = Connection('secret.json')
        self._conn.connect()

    def _initialise(self):
        self._mgr = Manager()
        self._mgr.load(self._conn)

    def run(self):
        if len(sys.argv) < 3:
            print('ERROR: term and section have not been set! ')
            return

        print('Connecting to OSM...')
        self._connect()
        self._initialise()

        self._set_term(sys.argv[1:3])
        if self._term is None:
            return

        print('Retrieving term programme...')
        programme = self._term.load_programme(self._conn)
        now = date.today()
        night = programme[0]
        for day in programme:
            night = day
            if day.date > now:
                break

        print('Retrieving members...')
        members = self._term.load_members(self._conn, include_data=True)

        print('Generating report...')
        template = ensureExtension(sys.argv[2]+'-Signin-Template', '.docx')
        filename = ensureExtension(sys.argv[2]+'-Signin', '.docx')
        document = Document(template)
        self._generate_report(programme, members, document, night)

        print('Saving to %s...' % (filename, ))
        document.save(filename)

        print('Done')

    def _set_term(self, args):
        term_name = args[0]
        self._set_section(args[1:])

        print('Setting term...')
        if term_name == 'current':
            term = self._section.current_term()
            if term is None:
                print('-> Currently not in a term')
                return
            else:
                self._term = term
                print('-> Term set to %s' % (str(term), ))
                return
        
        for term in self._section.terms:
            if term.name == term_name:
                self._term = term
                print('-> Term set to %s' % (str(term), ))
                return
            
        print('-> Unknown term: %s' % (term_name, ))

    def _set_section(self, args):
        print('Setting section...')
        section = self._mgr.find_section(args[0])
        if section is None:
            print('-> Unknown section: %s' % (args[0], ))
        else:
            self._section = section
            print('-> Section set to %s' % (str(section), ))
    
    def _generate_report(self, programme, members, doc, night):
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text == '{{date}}':
                    run.text = night.date.strftime('%d %B %Y')
                if run.text == '{{activity}}':
                    run.text = night.name

        print '...extracting data'
        report_data = []
        for person in members:
            if person.patrol == 'Leaders':
                continue

            name = '%s %s' % (person.first_name, person.last_name)
            try:
                contact = person.custom_data['contact_primary_1']
            except KeyError:
                try:
                    contact = person.custom_data['contact_primary_2']
                except KeyError:
                    contact = None

            if not contact is None:
                contact_person = '%s %s' % (contact['first_name'].strip(), contact['last_name'].strip())
                try:
                    contact_number = contact['mobile_phone']
                except KeyError:
                    try:
                        contact_number = contact['home_phone']
                    except KeyError:
                        contact_number = ''
            report_data.append([name, contact_person, contact_number])                 

        print '...populating table'
        report_data.sort(key=lambda r:r[0])
        row = 1
        table = doc.tables[0]
        for person in report_data:
            cells = table.rows[row].cells
            row += 1
            print('...adding row for %s...' % (person[0], ))
            cells[1].text = person[0]
            cells[4].text = person[1]
            cells[5].text = person[2]


def ensureExtension(filename, extension):
    return filename if filename.lower().endswith(extension) else filename + extension

def clearFormatting(paragraph, style=None):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if style is not None:
        paragraph.style = style

if __name__ == "__main__":
    mgr = ReportGenerator()
    mgr.run()
