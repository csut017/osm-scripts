'''
This script generates an audit report of which items each member has completed for the award scheme.
'''

import sys

from datetime import date
from docx import Document
from docx.shared import Cm
from osm import AwardScheme, Connection, Manager

class ReportGenerator(object):

    def __init__(self):
        self._conn = None
        self._mgr = None
        self._section = None
        self._term = None

    def _connect(self):
        self._conn = Connection('secret.json')
        self._conn.connect()

    def _initialise(self):
        self._mgr = Manager()
        self._mgr.load(self._conn)

    def run(self):
        print('Connecting to OSM...')
        self._connect()
        self._initialise()

        if len(sys.argv) < 3:
            print('ERROR: term and section have not been set! ')
            return

        self._set_term(sys.argv[1:3])
        if self._term is None:
            return

        print('Retrieving badge data...')
        scheme = AwardScheme(self._section.name + '-award.json')
        print('-> Loaded award scheme definition')
        self._term.load_badges(self._conn)
        badge_map = {}
        for badge in self._term.badges:
            badge_map[badge.badge_id] = badge
        print('-> Loaded badges')

        print('Generating report...')
        filename = ensureExtension(sys.argv[2]+'-Badge Audit', '.docx')
        document = Document()
        self._generate_header_footer(document)
        self._generate_report(scheme, document, badge_map)

        print('Saving to %s...' % (filename, ))
        document.save(filename)

        print('Done')

    def _set_term(self, args):
        term_name = args[0]
        self._set_section(args[1:])

        print('Setting term...')
        if term_name == 'current':
            term = self._section.current_term()
            if term is None:
                print('-> Currently not in a term')
                return
            else:
                self._term = term
                print('-> Term set to %s' % (str(term), ))
                return
        
        for term in self._section.terms:
            if term.name == term_name:
                self._term = term
                print('-> Term set to %s' % (str(term), ))
                return
            
        print('-> Unknown term: %s' % (term_name, ))

    def _set_section(self, args):
        print('Setting section...')
        section = self._mgr.find_section(args[0])
        if section is None:
            print('-> Unknown section: %s' % (args[0], ))
        else:
            self._section = section
            print('-> Section set to %s' % (str(section), ))
    
    def _generate_header_footer(self, document):
        now = date.today()
        section = document.sections[0]
        section.header.paragraphs[0].text = 'Badge Progress Report'
        section.footer.paragraphs[0].text = 'Generated ' + now.strftime('%d %B %Y')

    def _generate_report(self, scheme, document, badge_map):
        members = {}
        for badge in scheme.badges:
            for part in badge.parts:
                part.badge = badge_map[part.id]
                part.badge.load_progress(self._conn)
                part_map = dict((p.part_id, p.name.strip().lower()) for p in part.badge.parts)
                print('-> Loaded "%s"...' % (part.badge.name,))

                for person in part.badge.progress:
                    name = person.firstname + ' ' + person.lastname
                    for p_id, activity in person.parts.items():
                        p_name = part_map[p_id]

                        try:
                            items = members[name]
                        except KeyError:
                            items = {}
                            members[name] = items

                        try:
                            activities = items[p_name]                            
                        except KeyError:
                            activities = []
                            items[p_name] = activities
                        activities.append(activity + ' [' + badge.name + ']')
            
            print('-> Processed "%s"...' % (badge.name,))
        
        print('-> Generating final audit')
        for member in sorted(members.keys()):
            paragraph = document.add_paragraph(member.title())
            paragraph.style = document.styles['Heading 1']
            items = members[member]
            for item in sorted(items.keys()):
                paragraph = document.add_paragraph(item)
                paragraph.style = document.styles['Heading 2']
                for activity in sorted(items[item]):
                    document.add_paragraph(activity)
            print('-> Completed "%s"...' % (member,))

def ensureExtension(filename, extension):
    return filename if filename.lower().endswith(extension) else filename + extension

if __name__ == "__main__":
    mgr = ReportGenerator()
    mgr.run()
