'''
This script generates an overall progress report of where the division is at for the award scheme badges.
'''

import os
import sys

from datetime import date
from docx import Document
from docx.shared import Cm
from osm import AwardScheme, Connection, Manager

import matplotlib.pyplot as plt


class ReportGenerator(object):

    def __init__(self):
        self._conn = None
        self._mgr = None
        self._section = None
        self._term = None

    def _connect(self):
        self._conn = Connection('secret.json')
        self._conn.connect()

    def _initialise(self):
        self._mgr = Manager()
        self._mgr.load(self._conn)
        if not os.path.exists('temp_images'):
            os.makedirs('temp_images')

    def run(self):

        if len(sys.argv) < 3:
            print('ERROR: term and section have not been set! ')
            return

        print('Connecting to OSM...')
        self._connect()
        self._initialise()

        self._set_term(sys.argv[1:3])
        if self._term is None:
            return

        print('Retrieving badge data...')
        scheme = AwardScheme(self._section.name + '-award.json')
        print('-> Loaded award scheme definition')
        self._term.load_badges(self._conn)
        badge_map = {}
        for badge in self._term.badges:
            badge_map[badge.badge_id] = badge
        print('-> Loaded badges')

        print('Generating report...')
        filename = ensureExtension(sys.argv[2]+'-Badge Progress', '.docx')
        document = Document()
        self._generate_header_footer(document)
        self._generate_report(scheme, document, badge_map)

        print('Saving to %s...' % (filename, ))
        document.save(filename)

        print('Done')

    def _set_term(self, args):
        term_name = args[0]
        self._set_section(args[1:])

        print('Setting term...')
        if term_name == 'current':
            term = self._section.current_term()
            if term is None:
                print('-> Currently not in a term')
                return
            else:
                self._term = term
                print('-> Term set to %s' % (str(term), ))
                return
        
        for term in self._section.terms:
            if term.name == term_name:
                self._term = term
                print('-> Term set to %s' % (str(term), ))
                return
            
        print('-> Unknown term: %s' % (term_name, ))

    def _set_section(self, args):
        print('Setting section...')
        section = self._mgr.find_section(args[0])
        if section is None:
            print('-> Unknown section: %s' % (args[0], ))
        else:
            self._section = section
            print('-> Section set to %s' % (str(section), ))
    
    def _generate_header_footer(self, document):
        now = date.today()
        section = document.sections[0]
        section.header.paragraphs[0].text = 'Badge Progress Report'
        section.footer.paragraphs[0].text = 'Generated ' + now.strftime('%d %B %Y')

    def _generate_report(self, scheme, document, badge_map):
        for badge in scheme.badges:
            paragraph = document.add_paragraph(badge.name)
            paragraph.style = document.styles['Heading 1']
            counts, labels = [], []
            for part in badge.parts:
                progress = 0
                part.badge = badge_map[part.id]
                part.badge.load_progress(self._conn)
                print('-> Loaded "%s"...' % (part.badge.name,))
                for person in part.badge.progress:
                    progress += len(person.parts)

                mean = progress * 100 / len(part.badge.progress)
                completion = mean / len(part.badge.parts)
                labels.append(part.name)
                counts.append(completion)
            
            fig, ax = plt.subplots(1, 1, figsize=(5, len(badge.parts)))
            ax.set_xlim(0, 100)
            ax.set_xlabel('Percentage completed')
            ax.barh(labels, counts)
            badge_path = os.path.join('temp_images', badge.name + '.png')
            fig.savefig(badge_path, bbox_inches='tight', dpi=300)
            document.add_picture(badge_path, width=Cm(16))
            print('-> Generated "%s"...' % (badge.name,))


def ensureExtension(filename, extension):
    return filename if filename.lower().endswith(extension) else filename + extension

if __name__ == "__main__":
    mgr = ReportGenerator()
    mgr.run()
