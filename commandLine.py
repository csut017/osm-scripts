''' Script for downloading and working with a term programme. '''

import os
import sys
import traceback
from datetime import date

import xlsxwriter
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt
from osm import Connection, Manager


class ProgrammeManager(object):

    def __init__(self):
        self._conn = None
        self._mgr = None
        self._exit = False
        self._section = None
        self._term = None

        self._commands = {
            'q': self._exit_manager,
            'quit': self._exit_manager,
            'exit': self._exit_manager,
            'sections': self._list_sections,
            'section': self._set_section,
            'terms': self._list_terms,
            'term': self._set_term,
            'members': self._list_members,
            'programme': self._list_programme,
            'badges': self._list_badges,
            'badge': self._badge_actions,
        }

        self._badge_commands = {
            'parts': self._list_badge_parts,
            'progress': self._list_badge_progress,
            'export': self._export_badge_progress,
        }


    def _connect(self):
        self._conn = Connection('secret.json')
        self._conn.connect()


    def _initialise(self):
        self._mgr = Manager()
        self._mgr.load(self._conn)


    def run(self):
        print 'Connecting to OSM...'
        self._connect()
        self._initialise()

        print 'Welcome to OSM command line'
        if len(sys.argv) > 1:
            mgr._set_term(sys.argv[1:])

        while not self._exit:
            cmd_text = raw_input('>').strip()
            cmd_args = self._split_cmd(cmd_text)
            if len(cmd_args) == 0:
                continue

            cmd_name = cmd_args[0]
            cmd_args = cmd_args[1:]
            try:
                cmd = self._commands[cmd_name]
            except KeyError:
                print 'Unknown command: %s' % (cmd_name, )
                continue

            try:
                cmd(cmd_args)
            except Exception as ex:
                print 'Unexpected error: %s' % (str(ex), )
                traceback.print_exc()


    def _split_cmd(self, cmd_text):
        out = []
        word = ''
        in_string = False
        for c in cmd_text:
            if c == "'":
                in_string = not in_string
            elif c == ' ':
                if in_string:
                    word += c
                else:
                    if word != '':
                        out.append(word)
                    word = ''
            else:
                word += c

        if word != '':
            out.append(word)

        return out


    def _exit_manager(self, args):
        self._exit = True

    
    def _list_sections(self, args):
        for section in self._mgr.sections:
            print str(section)


    def _set_section(self, args):
        if len(args) < 1:
            print 'Current section is %s' % (str(self._section), )
            return

        section = self._mgr.find_section(args[0])
        if section is None:
            print 'Unknown section: %s' % (args[0], )
        else:
            self._section = section
            print 'Section set to %s' % (str(section), )

    
    def _list_terms(self, args):
        if self._section is None:
            print 'Section must be set first'
            return

        for term in self._section.terms:
            print str(term)


    def _set_term(self, args):
        if len(args) < 2 and self._section is None:
            print 'Section must be set first'
            return

        if len(args) < 1:
            print 'Current term is %s' % (str(self._term), )
            return

        term_name = args[0]
        if len(args) > 1:
            self._set_section(args[1:])

        if term_name == 'current':
            term = self._section.current_term()
            if term is None:
                print 'Currently not in a term'
                return
            else:
                self._term = term
                print 'Term set to %s' % (str(term), )
                return

        for term in self._section.terms:
            if term.name == term_name:
                self._term = term
                print 'Term set to %s' % (str(term), )
                return
            
        print 'Unknown term: %s' % (term_name, )

    def _list_members(self, args):
        if self._term is None:
            print 'Term must be set first'
            return

        if not self._term.members_loaded:
            print 'Loading members...'
            self._term.load_members(self._conn)
        for member in self._term.members:
            print str(member)

    def _list_programme(self, args):
        if self._term is None:
            print 'Term must be set first'
            return

        if not self._term.programme_loaded > 0:
            print 'Loading programme...'
            self._term.load_programme(self._conn)

        if len(args) >= 1:
            if args[0] == 'export':
                self._export_program(args)
            elif args[0] == 'members':
                if not self._term.programme_loaded > 1:
                    print '...loading attendance...'
                    self._term.load_programme(self._conn, True)
                for meeting in self._term.programme:
                    print str(meeting)
                    for member in meeting.members:
                        print '- ' + str(member)
            else:
                print 'Unknown command'
        else:
            for meeting in self._term.programme:
                print str(meeting)

    def _export_program(self, args):
        if len(args) < 2:
            print 'Missing filename'
            return

        filename = ensureExtension(args[1], '.xlsx')
        workbook = xlsxwriter.Workbook(filename)

        print '...exporting programme...'
        worksheet = workbook.add_worksheet(self._term.name)
        row = 1

        bold = workbook.add_format({'bold': True})
        worksheet.write('A1', 'Date', bold)
        worksheet.write('B1', 'Name', bold)
        worksheet.write('C1', 'Leader', bold)
        dateFormat = workbook.add_format({'num_format': 'd/m/yyyy'})
        for meeting in self._term.programme:
            worksheet.write_datetime(row, 0, meeting.date, dateFormat)
            worksheet.write(row, 1, meeting.name)
            worksheet.write(row, 2, meeting.leader)
            row += 1
        workbook.close()
        print '...done'

    def _list_badges(self, args):
        if self._term is None:
            print 'Term must be set first'
            return

        if len(args) >= 1:
            if args[0] == 'report':
                self._generate_badge_report(args[1:])
            elif args[0] == 'dump':
                self._dump_badges(args[1:])
            else:
                print 'Unknown command'
        else:
            if not self._term.badges_loaded:
                print 'Loading badges...'
                self._term.load_badges(self._conn)

            for badge in self._term.badges:
                print str(badge)

    def _dump_badges(self, args):
        if len(args) < 1:
            print 'Missing filename'
            return

        if not self._term.badges_loaded:
            print 'Loading badges...'
            self._term.load_badges(self._conn)

        print 'Dumping badges...'
        filename = ensureExtension(args[-1], '.xlsx')
        workbook = xlsxwriter.Workbook(filename)

        worksheet = workbook.add_worksheet('Badges')
        row = 1

        bold = workbook.add_format({'bold': True})
        worksheet.write('A1', 'Name', bold)
        worksheet.write('B1', 'Id', bold)
        worksheet.write('C1', 'Type', bold)
        worksheet.write('D1', 'Picture', bold)
        for badge in self._term.badges:
            worksheet.write(row, 0, badge.name)
            worksheet.write(row, 1, badge.badge_id)
            worksheet.write(row, 2, badge.type)
            worksheet.write(row, 3, badge.picture)
            row += 1
        workbook.close()

        print '...done'

    def _generate_badge_report(self, args):
        if len(args) < 1:
            print 'Missing filename'
            return

        print 'Retrieving badge report...'
        report = self._term.load_badges_by_person(self._conn)

        if not os.path.exists('badge_images'):
            os.makedirs('badge_images')

        print '...generating badge report...'
        now = date.today()
        filename = ensureExtension(args[-1], '.docx')
        document = Document()
        headingStyle = document.styles.add_style('TableHeading', WD_STYLE_TYPE.PARAGRAPH)
        headingStyle.font.bold=True

        section = document.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

        section.header.paragraphs[0].text = 'Badge Report'
        section.footer.paragraphs[0].text = 'Generated ' + now.strftime('%d %B %Y')
        table = document.add_table(rows = 1, cols = 2, style='Table Grid')
        table.columns[0].width = Cm(5)
        table.columns[1].width = Cm(21)
        cells = table.rows[0].cells
        cells[0].text = 'Person'
        cells[1].text = 'Badges'
        cells[0].width = Cm(5)
        cells[1].width = Cm(21)
        clearFormatting(cells[0].paragraphs[0], headingStyle)
        clearFormatting(cells[1].paragraphs[0], headingStyle)
        for person in report:
            cells = table.add_row().cells
            name = '%s %s' % (person.first_name, person.last_name)
            print '...adding row for %s...' % (name, )
            cells[0].text = name
            clearFormatting(cells[0].paragraphs[0])
            para = cells[1].paragraphs[0]
            clearFormatting(para)
            cells[0].width = Cm(5)
            cells[1].width = Cm(21)
            for badge in person.badges:
                badge_path = os.path.join('badge_images', os.path.basename(badge.picture))
                if not os.path.exists(badge_path):
                    print '...retrieving badge image for %s...' % (badge.name,)
                    self._conn.download_binary(badge.picture, badge_path)
                if badge.completed:
                    para.add_run().add_picture(badge_path, width = Cm(2))
                    para.add_run(' ')

        print '...saving...'
        document.save(filename)

        print '...done'

    def _badge_actions(self, args):
        if len(args) < 1:
            print 'Missing action'
            return

        if len(args) < 2:
            print 'Missing badge number'
            return

        if self._term is None:
            print 'Term must be set first'
            return

        if not self._term.badges_loaded:
            print 'Loading badges...'
            self._term.load_badges(self._conn)

        try:
            cmd = self._badge_commands[args[0]]
        except KeyError:
            print 'Unknown action %s' % (args[0], )
            return

        try:
            badge_number = int(args[1])            
            badge = self._term.badges[badge_number - 1]
        except IndexError:
            print 'Unknown badge %s' % (args[1], )
            return

        cmd(args[2:], badge)

    def _list_badge_parts(self, args, badge):
        print str(badge)
        for part in badge.parts:
            print str(part)

    def _list_badge_progress(self, args, badge):
        if not badge.progress_loaded:
            print 'Loading badge progress...'
            badge.load_progress(self._conn)

        for progress in badge.progress:
            print str(progress)

    def _export_badge_progress(self, args, badge):
        if len(args) < 1:
            print 'Missing filename'
            return

        filename = ensureExtension(args[-1], '.xlsx')
        badges = [self._term.badges[int(n) - 1] for n in args[:-1]]
        workbook = xlsxwriter.Workbook(filename)

        print 'Exporting badge progress...'
        print '...%s...' % (badge.name,)
        if not badge.progress_loaded:
            print '...loading badge progress...'
            badge.load_progress(self._conn)

        badge.export_progress(workbook=workbook)
        for other in badges:
            print '...%s...' % (other.name,)
            if not other.progress_loaded:
                print '...loading badge progress...'
                other.load_progress(self._conn)
            other.export_progress(workbook=workbook)
        workbook.close()
        print '...done'


def ensureExtension(filename, extension):
    return filename if filename.lower().endswith(extension) else filename + extension

def clearFormatting(paragraph, style=None):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if style is not None:
        paragraph.style = style

if __name__ == "__main__":
    mgr = ProgrammeManager()
    mgr.run()
